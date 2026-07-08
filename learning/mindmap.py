# -*- coding: utf-8 -*-
"""
Mind map — generate a memory-optimized mind map (Tony Buzan style) from the
transcript with AI, render it with markmap (colorful curved branches, native
zoom/pan), and let the user refine it by chatting.

AI returns Markdown (robust — no fragile syntax) + a JSON of thorough
per-node explanations. Export: HTML / PNG / PDF / DOCX / MD / TXT.
"""

import os
import re
import json
import threading
import datetime

from PySide6.QtCore import QObject, Signal, Qt, QTimer
from PySide6.QtCore import QUrl
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QComboBox, QLabel,
    QFileDialog, QLineEdit, QPlainTextEdit, QWidget,
)
from PySide6.QtWebEngineWidgets import QWebEngineView

from config import GEMINI_API_KEY, GEMINI_MODEL, SUPPORTED_LANGS


# ── markmap HTML shell (autoloader — colorful, zoom/pan, clickable) ───
_HTML = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  html,body {{ margin:0; height:100%; background:#ffffff; font-family:'Segoe UI',sans-serif; }}
  .markmap {{ position:absolute; inset:0; width:100%; height:100%; }}
  #tip {{ position:fixed; top:8px; left:50%; transform:translateX(-50%);
          color:#8a94a6; font-size:11px; pointer-events:none; z-index:5; }}
  #expl {{ position:fixed; left:14px; right:14px; bottom:14px; display:none; z-index:10;
           background:#1c2128; border:1px solid #30363d; border-radius:12px;
           padding:14px 18px; color:#e6edf3; font-size:14px; line-height:1.6;
           box-shadow:0 6px 24px rgba(0,0,0,.35); }}
  #expl b {{ color:#8957e5; font-size:15px; }}
  #expl .hint {{ color:#8a94a6; font-size:11px; float:right; }}
</style>
</head><body>
<div id="tip">scroll = zoom · drag = pan · click a node = explanation</div>
<div class="markmap">
<script type="text/template">
---
markmap:
  colorFreezeLevel: 2
  initialExpandLevel: -1
  maxWidth: 320
  spacingVertical: 10
  paddingX: 14
---
{md_block}
</script>
</div>
<div id="expl"></div>
<script>const EXPL = {expl_json}; const EDITABLE = {editable_js};</script>
<script src="https://cdn.jsdelivr.net/npm/markmap-autoloader@0.18"></script>
<script>
  function strip(s) {{ return (s||"").replace(/[*_`#]/g,"").replace(/\\s+/g," ").trim(); }}
  function lookup(k) {{
    k = strip(k);
    if (EXPL[k]) return EXPL[k];
    for (const key in EXPL) {{
      const sk = strip(key);
      if (sk && (sk.indexOf(k) !== -1 || k.indexOf(sk) !== -1)) return EXPL[key];
    }}
    return null;
  }}
  function attach() {{
    const nodes = document.querySelectorAll("g.markmap-node");
    if (!nodes.length) {{ setTimeout(attach, 500); return; }}
    nodes.forEach(function(n, idx) {{
      const fo = n.querySelector("foreignObject");
      if (!fo || fo.dataset.bound) return;
      fo.dataset.bound = "1";
      fo.style.cursor = "pointer";
      fo.addEventListener("click", function(ev) {{
        ev.stopPropagation();
        if (EDITABLE) {{
          document.title = "mmnode:" + idx + ":" + Date.now();
          return;
        }}
        const k = strip(fo.textContent);
        const e = lookup(k);
        const p = document.getElementById("expl");
        p.innerHTML = '<span class="hint">click anywhere to close</span><b>' + k +
                      '</b><br>' + (e ? e : "(no explanation)");
        p.style.display = "block";
      }});
    }});
    document.body.addEventListener("click", function(ev) {{
      if (!ev.target.closest("foreignObject") && !ev.target.closest("#expl"))
        document.getElementById("expl").style.display = "none";
    }});
  }}
  setTimeout(attach, 1200);
</script>
</body></html>"""


def build_html(markdown: str, explanations: dict | None = None,
               editable: bool = False) -> str:
    md = markdown.replace("</script", "<\\/script")
    return _HTML.format(
        md_block=md,
        expl_json=json.dumps(explanations or {}, ensure_ascii=False).replace("</", "<\\/"),
        editable_js="true" if editable else "false",
    )


# ── parsing helpers ───────────────────────────────────────────────────
def _clean_markdown(text: str) -> str:
    text = re.sub(r"```(?:markdown|md)?", "", text).strip()
    idx = text.find("#")
    if idx > 0:
        text = text[idx:]
    return text.strip()


def _strip_md(s: str) -> str:
    return re.sub(r"[*_`#]", "", s).strip()


def _parse_output(text: str) -> tuple[str, dict]:
    """Split AI output into (markdown, explanations_dict)."""
    code_part, _, expl_part = text.partition("---EXPLANATIONS---")
    md = _clean_markdown(code_part)
    expl = {}
    if expl_part.strip():
        m = re.search(r"\{.*\}", expl_part, flags=re.DOTALL)
        if m:
            try:
                raw = json.loads(m.group(0))
                expl = {_strip_md(str(k)): str(v) for k, v in raw.items()}
            except Exception:
                expl = {}
    return md, expl


def _has_tree(md: str) -> bool:
    """True if the markdown has a heading and at least one nested item."""
    has_h = bool(re.search(r"(?m)^#{1,6}\s+\S", md))
    has_item = bool(re.search(r"(?m)^\s*[-*+]\s+\S", md)) or \
        len(re.findall(r"(?m)^#{1,6}\s+\S", md)) >= 2
    return has_h and has_item


def parse_nodes(md: str) -> list[dict]:
    """Parse markdown into a flat node list: [{'level':int, 'text':raw}] (keeps **bold**)."""
    nodes = []
    cur_h = 0
    for raw in md.splitlines():
        s = raw.strip()
        if not s:
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)", s)
        if hm:
            cur_h = len(hm.group(1)) - 1
            nodes.append({"level": cur_h, "text": hm.group(2).strip()})
            continue
        bm = re.match(r"^[-*+]\s+(.*)", s)
        if bm:
            indent = len(raw) - len(raw.lstrip())
            nodes.append({"level": cur_h + 1 + indent // 2, "text": bm.group(1).strip()})
    if nodes:
        base = min(n["level"] for n in nodes)
        for n in nodes:
            n["level"] -= base
    return nodes


def nodes_to_md(nodes: list[dict]) -> str:
    """Serialize a node list back to markdown for markmap."""
    lines = []
    for n in nodes:
        lv, t = n["level"], n["text"]
        if lv == 0:
            lines.append(f"# {t}")
        elif lv == 1:
            lines.append(f"## {t}")
        else:
            lines.append("  " * (lv - 2) + f"- {t}")
    return "\n".join(lines)


def markdown_to_outline(md: str) -> list[tuple[int, str]]:
    """Parse markdown headings/bullets into [(level, label)]."""
    out = []
    cur_h = 0
    for raw in md.splitlines():
        s = raw.strip()
        if not s:
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)", s)
        if hm:
            cur_h = len(hm.group(1)) - 1
            out.append((cur_h, _strip_md(hm.group(2))))
            continue
        bm = re.match(r"^[-*+]\s+(.*)", s)
        if bm:
            indent = len(raw) - len(raw.lstrip())
            out.append((cur_h + 1 + indent // 2, _strip_md(bm.group(1))))
    if out:
        base = min(l for l, _ in out)
        out = [(l - base, t) for l, t in out]
    return out


# ── AI generation ─────────────────────────────────────────────────────
def _gen(prompt, qwen_getter):
    if GEMINI_API_KEY:
        from ai.gemini_client import get_gemini_client, fast_config
        client = get_gemini_client(GEMINI_API_KEY)
        cfg = fast_config(GEMINI_MODEL)
        kwargs = {"config": cfg} if cfg else {}
        r = client.models.generate_content(model=GEMINI_MODEL, contents=prompt, **kwargs)
        return _parse_output(getattr(r, "text", "") or "")
    qwen = qwen_getter()
    if not qwen:
        raise RuntimeError("No Gemini key and local Qwen unavailable")
    import torch
    msgs = [{"role": "user", "content": prompt}]
    p = qwen.tokenizer.apply_chat_template(
        msgs, tokenize=False, add_generation_prompt=True, enable_thinking=False)
    inputs = qwen.tokenizer(p, return_tensors="pt").to(qwen._device)
    with torch.no_grad():
        out = qwen.model.generate(**inputs, max_new_tokens=1400, do_sample=False,
                                  temperature=None, top_p=None)
    new = out[0][inputs["input_ids"].shape[1]:]
    raw = qwen.tokenizer.decode(new, skip_special_tokens=True)
    raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL)
    return _parse_output(raw)


_GEN_RULES = (
    "Design it for MEMORIZATION (Tony Buzan mind-map principles):\n"
    "- Output GitHub-flavored Markdown as a nested outline.\n"
    "- One single top-level heading '# ' = the central subject itself.\n"
    "- 3-6 second-level headings '## ' = the main branches (key themes).\n"
    "- Under each, nested bullets '- ' (and sub-bullets indented by 2 spaces) "
    "with only KEYWORDS, max ~5 words each — never full sentences.\n"
    "- Put the single most important word of each item in **bold** so the key "
    "point stands out.\n"
    "- Group related ideas so the structure itself aids recall.\n"
)


class MindmapWorker(QObject):
    done = Signal(str, dict)
    error = Signal(str)

    def __init__(self, text, lang_name, qwen_getter):
        super().__init__()
        self._text = text
        self._lang = lang_name
        self._qwen_getter = qwen_getter

    def start(self):
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        prompt = (
            f"Create a study mind map from the content below. "
            f"Write every label in {self._lang}.\n"
            f"{_GEN_RULES}"
            f"After the markdown, add a line '---EXPLANATIONS---' followed by a JSON "
            f"object mapping EVERY label (exact text, without markdown symbols) to a "
            f"THOROUGH 2-3 sentence explanation in {self._lang}. Each explanation must "
            f"make the concept clear AND give a concrete example or memory hook to help "
            f"remember it.\n"
            f"Reply with ONLY the markdown then the JSON. No code fences.\n\n"
            f"Content:\n{self._text}"
        )
        try:
            md, expl = _gen(prompt, self._qwen_getter)
            if not _has_tree(md):
                md, expl = _gen(
                    prompt + "\n\nYour previous attempt was invalid — you MUST use "
                             "'#' headings and indented '-' bullets.", self._qwen_getter)
            if not _has_tree(md):
                self.error.emit("AI did not return a valid mind map")
                return
            self.done.emit(md, expl)
        except Exception as e:
            self.error.emit(str(e))


class MindmapEditWorker(QObject):
    """Applies a conversational edit instruction to an existing mind map."""
    done = Signal(str, dict)
    error = Signal(str)

    def __init__(self, md, expl, instruction, qwen_getter):
        super().__init__()
        self._md = md
        self._expl = expl
        self._instr = instruction
        self._qwen_getter = qwen_getter

    def start(self):
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        prompt = (
            "You are refining an existing study mind map (Markdown outline) for a user, "
            "conversationally. Apply the user's request, keeping everything else "
            "unchanged and keeping the same language as the existing labels.\n"
            f"{_GEN_RULES}"
            "\nCurrent mind map (Markdown):\n"
            f"{self._md}\n\n"
            "Current explanations (JSON):\n"
            f"{json.dumps(self._expl, ensure_ascii=False)}\n\n"
            f"User request: {self._instr}\n\n"
            "Output the FULL updated Markdown, then a line '---EXPLANATIONS---' and the "
            "FULL updated JSON of thorough explanations (2-3 sentences each with a memory "
            "hook). No code fences."
        )
        try:
            md, expl = _gen(prompt, self._qwen_getter)
            if not _has_tree(md):
                self.error.emit("AI did not return a valid mind map")
                return
            merged = dict(self._expl)
            merged.update(expl or {})
            self.done.emit(md, merged)
        except Exception as e:
            self.error.emit(str(e))


MINDMAP_QSS = """
QDialog { background:#0d1117; color:#c9d1d9; }
QComboBox { background:#1c2128; border:1px solid #30363d; border-radius:6px; padding:4px 8px; color:#c9d1d9; }
QComboBox QAbstractItemView { background:#161b22; color:#c9d1d9; border:1px solid #30363d; selection-background-color:#1f6feb; }
QPushButton { background:#21262d; border:1px solid #30363d; border-radius:6px; padding:5px 12px; color:#c9d1d9; }
QPushButton:hover { background:#30363d; }
QPushButton#send { background:#1f6feb; border:none; color:white; font-weight:500; }
QPushButton#send:hover { background:#2b7bf3; }
QLineEdit { background:#0a0c10; border:1px solid #30363d; border-radius:8px; padding:7px 10px; color:#e6edf3; }
QLineEdit:focus { border:1px solid #58a6ff; }
QLabel { color:#8b949e; }
"""


class MindmapDialog(QDialog):
    def __init__(self, parent, markdown, explanations=None, default_dir=None,
                 qwen_getter=None, lang_code="en"):
        super().__init__(parent)
        self._nodes = parse_nodes(markdown)
        self._md = nodes_to_md(self._nodes)
        self._expl = explanations or {}
        self._dir = default_dir or ""
        self._qwen_getter = qwen_getter or (lambda: None)
        self._lang_code = lang_code
        self._lang_codes = list(SUPPORTED_LANGS.keys())
        self._history = []
        self._edit_worker = None
        self._sel = None

        self.setWindowTitle("🗺 Mind map")
        self.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        self.resize(900, 680)
        self.setStyleSheet(MINDMAP_QSS)

        lay = QVBoxLayout(self)
        lay.setContentsMargins(10, 10, 10, 10)
        lay.setSpacing(8)

        # ── Top bar: undo · language · export ──
        bar = QHBoxLayout(); bar.setSpacing(6)
        self._undo_btn = QPushButton("↩ Undo")
        self._undo_btn.setEnabled(False)
        self._undo_btn.clicked.connect(self._undo)
        bar.addWidget(self._undo_btn)
        bar.addStretch()
        bar.addWidget(QLabel("Language"))
        self._lang_combo = QComboBox()
        self._lang_combo.addItems(
            [f"{SUPPORTED_LANGS[c]['flag']} {SUPPORTED_LANGS[c]['name']}" for c in self._lang_codes])
        if lang_code in self._lang_codes:
            self._lang_combo.setCurrentIndex(self._lang_codes.index(lang_code))
        self._lang_combo.currentIndexChanged.connect(self._change_lang)
        bar.addWidget(self._lang_combo)
        bar.addWidget(QLabel("  Export"))
        self._fmt = QComboBox()
        self._fmt.addItems(["html", "png", "pdf", "docx", "md", "txt"])
        bar.addWidget(self._fmt)
        b_save = QPushButton("💾 Save")
        b_save.clicked.connect(self._export)
        bar.addWidget(b_save)
        lay.addLayout(bar)

        self.view = QWebEngineView()
        self.view.titleChanged.connect(self._on_title)
        lay.addWidget(self.view, 1)

        # ── Direct node editor (hidden until a node is clicked) ──
        self._edit_panel = QWidget()
        ep = QVBoxLayout(self._edit_panel)
        ep.setContentsMargins(0, 0, 0, 0); ep.setSpacing(4)
        r1 = QHBoxLayout(); r1.setSpacing(6)
        r1.addWidget(QLabel("✏ Node"))
        self._e_label = QLineEdit()
        self._e_label.setPlaceholderText("Node label (use **word** for bold)")
        r1.addWidget(self._e_label, 1)
        b_savenode = QPushButton("Save"); b_savenode.setObjectName("send")
        b_savenode.clicked.connect(self._save_node); r1.addWidget(b_savenode)
        b_child = QPushButton("＋ Child"); b_child.clicked.connect(self._add_child); r1.addWidget(b_child)
        b_del = QPushButton("🗑"); b_del.setFixedWidth(36); b_del.clicked.connect(self._del_node); r1.addWidget(b_del)
        b_close = QPushButton("✕"); b_close.setFixedWidth(30)
        b_close.clicked.connect(lambda: self._edit_panel.setVisible(False)); r1.addWidget(b_close)
        ep.addLayout(r1)
        self._e_expl = QPlainTextEdit()
        self._e_expl.setPlaceholderText("Explanation for this node…")
        self._e_expl.setFixedHeight(54)
        ep.addWidget(self._e_expl)
        self._edit_panel.setVisible(False)
        lay.addWidget(self._edit_panel)

        # ── Conversational (AI) edit ──
        edit_row = QHBoxLayout(); edit_row.setSpacing(6)
        self._prompt = QLineEdit()
        self._prompt.setPlaceholderText(
            "Or refine with AI…  e.g. \"add a branch about costs\", "
            "\"explain the progress branch in more detail\"")
        self._prompt.returnPressed.connect(self._apply_edit)
        edit_row.addWidget(self._prompt, 1)
        self._edit_btn = QPushButton("Send")
        self._edit_btn.setObjectName("send")
        self._edit_btn.setFixedWidth(72)
        self._edit_btn.clicked.connect(self._apply_edit)
        edit_row.addWidget(self._edit_btn)
        lay.addLayout(edit_row)

        self._status = QLabel("")
        lay.addWidget(self._status)

        self._render()

    # ── Rendering / undo ──────────────────────────────────────────────
    def _render(self):
        self._md = nodes_to_md(self._nodes)
        self.view.setHtml(build_html(self._md, self._expl, editable=True))

    def _push_undo(self):
        self._history.append(([dict(n) for n in self._nodes], dict(self._expl)))
        self._undo_btn.setEnabled(True)

    def _undo(self):
        if not self._history:
            return
        self._nodes, self._expl = self._history.pop()
        self._undo_btn.setEnabled(bool(self._history))
        self._sel = None
        self._edit_panel.setVisible(False)
        self._render()
        self._flash("↩ Reverted")

    # ── Direct node editing ───────────────────────────────────────────
    def _on_title(self, title):
        if not title.startswith("mmnode:"):
            return
        try:
            idx = int(title.split(":")[1])
        except Exception:
            return
        if not (0 <= idx < len(self._nodes)):
            return
        self._sel = idx
        n = self._nodes[idx]
        self._e_label.setText(n["text"])
        self._e_expl.setPlainText(self._expl_for(_strip_md(n["text"])) or "")
        self._edit_panel.setVisible(True)
        self._e_label.setFocus()

    def _save_node(self):
        if self._sel is None:
            return
        self._push_undo()
        n = self._nodes[self._sel]
        old_key = _strip_md(n["text"])
        n["text"] = self._e_label.text().strip() or n["text"]
        new_key = _strip_md(n["text"])
        e = self._e_expl.toPlainText().strip()
        if old_key != new_key:
            self._expl.pop(old_key, None)
        if e:
            self._expl[new_key] = e
        else:
            self._expl.pop(new_key, None)
        self._render()
        self._flash("✅ Node saved")

    def _add_child(self):
        if self._sel is None:
            return
        self._push_undo()
        lvl = self._nodes[self._sel]["level"] + 1
        self._nodes.insert(self._sel + 1, {"level": lvl, "text": "New node"})
        self._sel += 1
        self._e_label.setText("New node")
        self._e_expl.setPlainText("")
        self._render()
        self._flash("➕ Child added — edit label then Save")
        self._e_label.setFocus(); self._e_label.selectAll()

    def _del_node(self):
        if self._sel is None:
            return
        i = self._sel
        lvl = self._nodes[i]["level"]
        j = i + 1
        while j < len(self._nodes) and self._nodes[j]["level"] > lvl:
            j += 1
        self._push_undo()
        for n in self._nodes[i:j]:
            self._expl.pop(_strip_md(n["text"]), None)
        del self._nodes[i:j]
        if not self._nodes:
            self._nodes = [{"level": 0, "text": "Mind map"}]
        self._sel = None
        self._edit_panel.setVisible(False)
        self._render()
        self._flash("🗑 Node deleted")

    # ── Language ──────────────────────────────────────────────────────
    def _change_lang(self):
        code = self._lang_codes[self._lang_combo.currentIndex()]
        if code == self._lang_code or self._edit_worker is not None:
            return
        self._lang_code = code
        name = SUPPORTED_LANGS.get(code, {}).get("name", code)
        self._run_ai_edit(
            f"Translate ALL node labels and ALL explanations to {name}. "
            f"Keep the exact same structure and hierarchy.")

    # ── Conversational (AI) editing ───────────────────────────────────
    def _apply_edit(self):
        instr = self._prompt.text().strip()
        if not instr:
            return
        self._prompt.clear()
        self._run_ai_edit(instr)

    def _run_ai_edit(self, instr):
        if self._edit_worker is not None:
            return
        self._edit_btn.setEnabled(False)
        self._edit_btn.setText("…")
        self._status.setText(f"🤖 {instr[:60]}")
        self._edit_worker = MindmapEditWorker(
            self._md, self._expl, instr, self._qwen_getter)
        self._edit_worker.done.connect(self._on_edit_done)
        self._edit_worker.error.connect(self._on_edit_error)
        self._edit_worker.start()

    def _on_edit_done(self, md, expl):
        self._push_undo()
        self._nodes = parse_nodes(md)
        self._expl = expl
        self._sel = None
        self._edit_panel.setVisible(False)
        self._render()
        self._edit_worker = None
        self._edit_btn.setEnabled(True)
        self._edit_btn.setText("Send")
        self._flash("✅ Map updated")

    def _on_edit_error(self, msg):
        self._edit_worker = None
        self._edit_btn.setEnabled(True)
        self._edit_btn.setText("Send")
        self._flash(f"⚠️ Edit failed: {msg}")

    # ── Export ────────────────────────────────────────────────────────
    def _export(self):
        fmt = self._fmt.currentText()
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        default = os.path.join(self._dir, f"mindmap_{ts}.{fmt}")
        path, _ = QFileDialog.getSaveFileName(self, "Save mind map", default, f"*.{fmt}")
        if not path:
            return
        try:
            if fmt == "html":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(build_html(self._md, self._expl))
            elif fmt == "png":
                self.view.grab().save(path, "PNG")
            elif fmt == "pdf":
                self.view.page().printToPdf(path)
            elif fmt == "docx":
                self._export_docx(path)
            elif fmt in ("md", "txt"):
                self._export_outline(path, fmt)
            self._flash(f"💾 Saved: {os.path.basename(path)}")
        except Exception as e:
            self._flash(f"⚠️ Export error: {e}")

    def _expl_for(self, label):
        if label in self._expl:
            return self._expl[label]
        for k, v in self._expl.items():
            if k in label or label in k:
                return v
        return None

    def _export_outline(self, path, fmt):
        lines = []
        for level, label in markdown_to_outline(self._md):
            e = self._expl_for(label)
            if fmt == "md":
                lines.append(("#" * min(level + 1, 6) + " " + label) if level == 0
                             else ("  " * (level - 1) + "- **" + label + "**"))
                if e:
                    lines.append(("  " * max(level - 1, 0)) + "  " + e)
            else:
                lines.append("  " * level + ("• " if level else "") + label)
                if e:
                    lines.append("  " * (level + 1) + "→ " + e)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")

    def _export_docx(self, path):
        from docx import Document
        doc = Document()
        for level, label in markdown_to_outline(self._md):
            if level == 0:
                doc.add_heading(label, level=0)
            elif level == 1:
                doc.add_heading(label, level=1)
            else:
                doc.add_paragraph(label, style="List Bullet")
            e = self._expl_for(label)
            if e:
                doc.add_paragraph(e)
        doc.save(path)

    def _flash(self, msg):
        self._status.setText(msg)
        QTimer.singleShot(3000, lambda: self._status.setText(""))
